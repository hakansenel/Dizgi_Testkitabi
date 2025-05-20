import json
import os
import glob
import re
import base64
import traceback
from io import BytesIO
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from PIL import Image
from bs4 import BeautifulSoup
# Dosya yolları - config.py'den al
from config import JSON_INPUT_DIR, JSON_OUTPUT_DIR, WORD_OUTPUT_DIR


# Klasör yollarını ayarla
input_dir = ("C:/Users/Öykü/Desktop/claude/solorota/docs/JSON")
json_output_dir = ("C:/Users/Öykü/Desktop/claude/solorota/docs/JSON_processed")
log_file = os.path.join(json_output_dir, "conversion_log.txt")

# Klasörleri oluştur (yoksa)
for directory in [json_output_dir]: # type: ignore
    if not os.path.exists(directory):
        os.makedirs(directory)

# Standart alan yapısı ve varsayılan değerler
standard_format = {
    "soruYazari": "",
    "sinifDuzeyi": "",
    "ders": "",
    "zorluk": 0,
    "kazanim": [],
    "konu": "",
    "unite": "",
    "ustMetin": "",
    "gorsel": "",
    "soruVideosuLinki": None,
    "soruMetni": "",
    "secenekler": {},
    "dogruCevap": "",
    "cozum": "",
    "cozumVideosuLinki": None
}

# Hata ve işlem kaydı tutucu
class Logger:
    def __init__(self, log_file):
        self.errors = []
        self.warnings = []
        self.info = []
        self.log_file = log_file
        
    def log_error(self, file_name, error_type, error_message, soru_index=None):
        timestamp = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        error_entry = {
            'time': timestamp,
            'file': file_name,
            'type': error_type,
            'message': error_message,
            'soru_index': soru_index
        }
        self.errors.append(error_entry)
        index_info = f", Soru: {soru_index}" if soru_index is not None else ""
        print(f"HATA: {error_type} - {error_message} (Dosya: {file_name}{index_info})")
        
    def log_warning(self, file_name, warning_type, warning_message, soru_index=None):
        timestamp = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        warning_entry = {
            'time': timestamp,
            'file': file_name,
            'type': warning_type,
            'message': warning_message,
            'soru_index': soru_index
        }
        self.warnings.append(warning_entry)
        index_info = f", Soru: {soru_index}" if soru_index is not None else ""
        print(f"UYARI: {warning_type} - {warning_message} (Dosya: {file_name}{index_info})")
        
    def log_info(self, file_name, info_message):
        timestamp = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        info_entry = {
            'time': timestamp,
            'file': file_name,
            'message': info_message
        }
        self.info.append(info_entry)
        
    def save_log(self):
        with open(self.log_file, 'w', encoding='utf-8') as f:
            f.write("JSON DÖNÜŞÜM RAPORU\n")
            f.write("=" * 80 + "\n\n")
            
            # Hataları yaz
            if self.errors:
                f.write("HATALAR\n")
                f.write("-" * 80 + "\n")
                for error in self.errors:
                    f.write(f"Zaman: {error['time']}\n")
                    f.write(f"Dosya: {error['file']}\n")
                    if error['soru_index'] is not None:
                        f.write(f"Soru Index: {error['soru_index']}\n")
                    f.write(f"Hata Tipi: {error['type']}\n")
                    f.write(f"Mesaj: {error['message']}\n")
                    f.write("-" * 80 + "\n\n")
            
            # Uyarıları yaz
            if self.warnings:
                f.write("UYARILAR\n")
                f.write("-" * 80 + "\n")
                for warning in self.warnings:
                    f.write(f"Zaman: {warning['time']}\n")
                    f.write(f"Dosya: {warning['file']}\n")
                    if warning['soru_index'] is not None:
                        f.write(f"Soru Index: {warning['soru_index']}\n")
                    f.write(f"Uyarı Tipi: {warning['type']}\n")
                    f.write(f"Mesaj: {warning['message']}\n")
                    f.write("-" * 80 + "\n\n")
            
            # Bilgileri yaz
            if self.info:
                f.write("BİLGİLER\n")
                f.write("-" * 80 + "\n")
                for info in self.info:
                    f.write(f"Zaman: {info['time']}\n")
                    f.write(f"Dosya: {info['file']}\n")
                    f.write(f"Mesaj: {info['message']}\n")
                    f.write("-" * 80 + "\n\n")
            
            print(f"\nDönüşüm raporu kaydedildi: {self.log_file}")

# HTML tablo algılama fonksiyonu
def contains_html_table(text):
    """HTML tablosu içerip içermediğini kontrol eder"""
    if not text or not isinstance(text, str):
        return False
    
    return '<table' in text.lower() and '</table>' in text.lower()

# HTML tablosunu Word tablosuna dönüştüren fonksiyon
def html_table_to_word(doc, html_table):
    """HTML tablosunu Word tablosuna dönüştürür"""
    try:
        # BeautifulSoup ile tabloyu ayrıştır
        soup = BeautifulSoup(html_table, 'html.parser')
        table_elem = soup.find('table')
        
        if not table_elem:
            return False
        
        # Satır ve sütunları say
        rows = table_elem.find_all('tr')
        if not rows:
            return False
            
        # En fazla sütun sayısını bul
        max_cols = 0
        for row in rows:
            cols = row.find_all(['th', 'td'])
            max_cols = max(max_cols, len(cols))
        
        if max_cols == 0:
            return False
        
        # Word tablosu oluştur
        word_table = doc.add_table(rows=len(rows), cols=max_cols)
        word_table.style = 'Table Grid'
        
        # Tabloyu doldur
        for i, row in enumerate(rows):
            cells = row.find_all(['th', 'td'])
            
            for j, cell in enumerate(cells):
                if j < max_cols:  # Sütun sayısı sınırını aşmayı önle
                    # Başlık hücresi mi kontrol et (th)
                    is_header = cell.name == 'th'
                    
                    # Hücre metnini al
                    cell_text = cell.get_text(strip=True)
                    
                    # Word tablosuna ekle
                    word_cell = word_table.cell(i, j)
                    
                    # Hücre içeriğine stil uygula
                    paragraph = word_cell.paragraphs[0]
                    run = paragraph.add_run(cell_text)
                    
                    # Başlık hücresi ise kalın yap
                    if is_header:
                        run.bold = True
                        # Başlık hücresi arkaplan rengini ayarla
                        shade_cell(word_cell, "DDDDDD")  # Açık gri
        
        # Tabloyu otomatik genişleştir
        word_table.autofit = True
        return True
        
    except Exception as e:
        print(f"HTML tablosu dönüştürme hatası: {str(e)}")
        return False

# Hücre arkaplan rengini ayarlama
def shade_cell(cell, color):
    """Word tablosundaki hücreye arkaplan rengi ekler"""
    try:
        shading = OxmlElement('w:shd')
        shading.set(qn('w:fill'), color)
        cell._tc.get_or_add_tcPr().append(shading)
    except:
        pass  # Hücre rengini ayarlayamazsa devam et

# HTML etiketlerini tamamen temizleyen fonksiyon
def remove_html_tags(text):
    """HTML etiketlerini metinden tamamen temizler"""
    if not text or not isinstance(text, str):
        return text
    
    # HTML etiketlerini tamamen temizle
    pattern = r'<.*?>'
    cleaned_text = re.sub(pattern, '', text)
    
    # HTML karakter referanslarını dönüştür
    html_entities = {
        '&nbsp;': ' ', 
        '&lt;': '<', 
        '&gt;': '>', 
        '&amp;': '&', 
        '&quot;': '"', 
        '&apos;': "'",
        '&ldquo;': '"',
        '&rdquo;': '"',
        '&lsquo;': ''',
        '&rsquo;': ''',
        '&ndash;': '–',
        '&mdash;': '—',
        '&hellip;': '…'
    }
    
    for entity, char in html_entities.items():
        cleaned_text = cleaned_text.replace(entity, char)
    
    return cleaned_text

# Base64 görsel kontrolü
def is_base64_image(text):
    """Metnin base64 formatında bir görsel olup olmadığını kontrol eder"""
    if not text or not isinstance(text, str):
        return False
    
    # Base64 formatı kontrolü
    base64_patterns = [
        r'^data:image\/[a-zA-Z0-9-]+;base64,',  # HTML5 data URI formatı
        r'^[A-Za-z0-9+\/=]+$'  # Sadece base64 karakterleri içeren string
    ]
    
    # Base64 formatı kontrol et
    is_base64 = False
    for pattern in base64_patterns:
        if re.match(pattern, text):
            is_base64 = True
            break
    
    # Base64 formatında değilse veya çok kısa ise (muhtemelen görsel değil)
    if not is_base64 or len(text) < 100:
        return False
    
    # Geçerli base64 verisi mi kontrol et
    try:
        if ';base64,' in text:
            _, base64_data = text.split(';base64,', 1)
        else:
            base64_data = text
            
        # Base64 karakterlerini düzelt
        base64_data = base64_data.replace(' ', '+')
        
        # Padding düzelt
        missing_padding = len(base64_data) % 4
        if missing_padding:
            base64_data += '=' * (4 - missing_padding)
            
        # Base64'ü çözmeyi dene
        base64.b64decode(base64_data)
        return True
    except:
        return False

# Base64 görseli işleme
def process_base64_image(base64_string):
    """Base64 formatındaki görsel verisini işler"""
    if not base64_string or not isinstance(base64_string, str):
        return None
    
    try:
        # Base64 prefix'i varsa ayır
        if ';base64,' in base64_string:
            _, base64_data = base64_string.split(';base64,', 1)
        else:
            base64_data = base64_string
        
        # Base64 karakterlerini düzelt
        base64_data = base64_data.replace(' ', '+')
        
        # Padding düzelt
        missing_padding = len(base64_data) % 4
        if missing_padding:
            base64_data += '=' * (4 - missing_padding)
        
        # Base64'ü çöz
        image_data = base64.b64decode(base64_data)
        return image_data
    except Exception as e:
        print(f"Base64 çözme hatası: {str(e)}")
        return None

# Word belgesine görsel ekleme
def add_image_to_document(doc, image_data, width=None):
    """Görsel verisini Word belgesine ekler"""
    if not image_data:
        return False
        
    try:
        image_stream = BytesIO(image_data)
        
        # Word'e ekle
        if width:
            doc.add_picture(image_stream, width=Inches(width))
        else:
            # Varsayılan genişlik
            doc.add_picture(image_stream, width=Inches(5))
        
        return True
    except Exception as e:
        print(f"Resim ekleme hatası: {str(e)}")
        return False

# HTML metni içindeki tabloları ve metin bölümlerini ayırır
def split_html_text_and_tables(html_text):
    """HTML metni içindeki tabloları ve metin bölümlerini ayırır"""
    if not html_text or not isinstance(html_text, str):
        return []
    
    # BeautifulSoup ile HTML'i ayrıştır
    soup = BeautifulSoup(html_text, 'html.parser')
    
    # Sonuç parçalarını tut
    parts = []
    
    # İlk metin parçasını al (tabloya kadar)
    current_pos = 0
    
    # Tüm tabloları bul ve metinle ayır
    for table in soup.find_all('table'):
        # Tablonun HTML'deki pozisyonunu bul
        table_html = str(table)
        table_start = html_text.find(table_html, current_pos)
        
        if table_start > current_pos:
            # Tablo öncesi metin
            text_before = html_text[current_pos:table_start]
            if text_before.strip():
                parts.append(('text', text_before))
        
        # Tablo
        parts.append(('table', table_html))
        
        # Pozisyonu güncelle
        current_pos = table_start + len(table_html)
    
    # Son metin parçası
    if current_pos < len(html_text):
        text_after = html_text[current_pos:]
        if text_after.strip():
            parts.append(('text', text_after))
    
    return parts

# HTML etiketli metinleri Word'de formatlı olarak gösterme
def apply_html_formatting_to_word(paragraph, text):
    """HTML formatlarını Word'e uygular (tekrarlama olmadan)"""
    if not text or not isinstance(text, str):
        return
    
    # HTML etiketleri var mı kontrol et
    if not re.search(r'<[bui]>', text) and not re.search(r'<table', text):
        # HTML etiketi yoksa direkt ekle
        paragraph.add_run(remove_html_tags(text))
        return
    
    # HTML formatlarını belirle ve pozisyonları kaydet
    # Altı çizili (<u>), kalın (<b>) ve italik (<i>) etiketleri
    format_positions = []
    
    # <u> etiketleri
    for match in re.finditer(r'<u>(.*?)</u>', text, re.DOTALL):
        format_positions.append({
            'type': 'underline',
            'start': match.start(),
            'end': match.end(),
            'content': match.group(1)
        })
    
    # <b> etiketleri
    for match in re.finditer(r'<b>(.*?)</b>', text, re.DOTALL):
        format_positions.append({
            'type': 'bold',
            'start': match.start(),
            'end': match.end(),
            'content': match.group(1)
        })
    
    # <i> etiketleri
    for match in re.finditer(r'<i>(.*?)</i>', text, re.DOTALL):
        format_positions.append({
            'type': 'italic',
            'start': match.start(),
            'end': match.end(),
            'content': match.group(1)
        })
    
    # Format yoksa düz metni ekle
    if not format_positions:
        paragraph.add_run(remove_html_tags(text))
        return
    
    # Temiz metin oluştur (HTML etiketleri olmadan)
    clean_text = remove_html_tags(text)
    
    # HTML etiketlerini içeren metnin işlenmiş parçaları
    # Örnek: <u>altı çizili</u> normal <b>kalın</b>
    # Parçalar: ["altı çizili" (underline), " normal ", "kalın" (bold)]
    
    # Öncelikle her bir etiketin temiz metindeki pozisyonunu belirle
    clean_format_positions = []
    for fmt in format_positions:
        content = remove_html_tags(fmt['content'])
        start_idx = clean_text.find(content)
        if start_idx >= 0:
            clean_format_positions.append({
                'type': fmt['type'],
                'start': start_idx,
                'end': start_idx + len(content),
                'content': content
            })
    
    # Temiz metni bölümle
    segments = []
    current_pos = 0
    
    # Tüm metin pozisyonlarını belirle
    break_points = set()
    for fmt in clean_format_positions:
        break_points.add(fmt['start'])
        break_points.add(fmt['end'])
    
    break_points = sorted(break_points)
    
    # Bölümleri oluştur
    for pos in break_points:
        if pos > current_pos:
            # Bu bölüm için etiketleri belirle
            active_formats = [fmt['type'] for fmt in clean_format_positions 
                             if fmt['start'] <= current_pos < fmt['end']]
            
            segments.append({
                'text': clean_text[current_pos:pos],
                'formats': active_formats
            })
            
            current_pos = pos
    
    # Son bölüm
    if current_pos < len(clean_text):
        active_formats = [fmt['type'] for fmt in clean_format_positions 
                         if fmt['start'] <= current_pos < fmt['end']]
        
        segments.append({
            'text': clean_text[current_pos:],
            'formats': active_formats
        })
    
    # Bölümleri Word'e ekle
    for segment in segments:
        run = paragraph.add_run(segment['text'])
        
        # Formatları uygula
        if 'bold' in segment['formats']:
            run.bold = True
        if 'underline' in segment['formats']:
            run.underline = True
        if 'italic' in segment['formats']:
            run.italic = True

# HTML metinleri ve tabloları işleme ve Word'e ekleme
def process_html_content_to_word(doc, html_content, file_name, logger, soru_index=None):
    """HTML içeriğini (metin ve tablolar dahil) Word belgesine ekler"""
    if not html_content or not isinstance(html_content, str):
        return
    
    # Tablo içeriyor mu kontrol et
    if contains_html_table(html_content):
        try:
            # HTML'i metin ve tablo parçalarına ayır
            parts = split_html_text_and_tables(html_content)
            
            # Her bir parçayı işle
            for part_type, part_content in parts:
                if part_type == 'table':
                    # Tablo ekle
                    if not html_table_to_word(doc, part_content):
                        logger.log_warning(file_name, "Tablo Dönüştürme", "Tablo işlenirken hata oluştu", soru_index)
                        # Tablo dönüştürülemezse, metin olarak ekle
                        p = doc.add_paragraph()
                        p.add_run(remove_html_tags(part_content))
                else:
                    # Metin ekle
                    if part_content.strip():
                        p = doc.add_paragraph()
                        apply_html_formatting_to_word(p, part_content)
        except Exception as e:
            logger.log_error(file_name, "HTML İşleme", f"HTML içeriği işlenirken hata: {str(e)}", soru_index)
            # Hata durumunda düz metin olarak ekle
            p = doc.add_paragraph()
            p.add_run(remove_html_tags(html_content))
    else:
        # Tablo içermiyor, normal HTML formatlaması uygula
        p = doc.add_paragraph()
        apply_html_formatting_to_word(p, html_content)

# JSON içerik düzeltme
def fix_json_content(content):
    """JSON içeriğinde yaygın hataları kontrol eder"""
    common_issues = [
        ('"..."', '"..."'),  # Yanlış tırnak karakteri
        ('…', '...'),        # Yanlış üç nokta
        ('\'', "'"),         # Yanlış tek tırnak
        ('"', '"'),          # Yanlış çift tırnak
        ('–', '-'),          # Yanlış tire
        ('\n]', '\n  ]'),    # Yanlış girintileme
        ('\n}', '\n  }'),    # Yanlış girintileme
    ]
    
    fixed_content = content
    for old, new in common_issues:
        if old in fixed_content:
            fixed_content = fixed_content.replace(old, new)
    
    # Görsel özel düzeltme - JSON ayrıştırma öncesinde karmaşık yapıyı basit yapıya dönüştür
    pattern = r'"gorsel"\s*:\s*{\s*"image"\s*:\s*{\s*"mime"\s*:\s*"[^"]+"\s*,\s*"data"\s*:\s*"([^"]+)"\s*}\s*}'
    fixed_content = re.sub(pattern, r'"gorsel":"\1"', fixed_content)
            
    return fixed_content

# Kompleks JSON yapısından soru listesi çıkarma
def extract_questions(data, file_name, logger):
    """
    Farklı JSON yapılarından soru listesi çıkarır
    Format örnekleri:
    1. {"sorular": [{soru1}, {soru2}]}
    2. [{soru1}, {soru2}]
    3. {"kolay": {soru1}, "orta": {soru2}}
    4. [{"soru": {soru1}}, {"soru": {soru2}}]
    5. [{"paralelSorular": {"soru1_paralel1": {soru1}, "soru1_paralel2": {soru2}}}]
    6. [{ "sorular": [{soru1}, {soru2}]}]
    """
    questions = []
    
    def is_valid_question(obj):
        """Bir objenin geçerli soru formatında olup olmadığını kontrol eder"""
        required_minimal_fields = ['soruMetni', 'secenekler', 'dogruCevap']
        
        # Soru nesnesi kontrolü ve gerekli alanların varlığı
        is_valid = isinstance(obj, dict) and all(field in obj for field in required_minimal_fields)
        
        # Gerekli kontroller
        if is_valid:
            # soruMetni string olmalı veya dönüştürülebilir olmalı
            if not isinstance(obj['soruMetni'], str) and obj['soruMetni'] is not None:
                obj['soruMetni'] = str(obj['soruMetni'])
                
            # secenekler dictionary olmalı
            if not isinstance(obj['secenekler'], dict):
                return False
                
            # dogruCevap string olmalı veya dönüştürülebilir olmalı
            if not isinstance(obj['dogruCevap'], str) and obj['dogruCevap'] is not None:
                obj['dogruCevap'] = str(obj['dogruCevap'])
        
        return is_valid
    
    def process_dict(d, path=""):
        """Dictionary tipindeki veriyi işler"""
        # Eğer direkt soru objesi ise
        if is_valid_question(d):
            questions.append((d, "root"))
            return
            
        for key, value in d.items():
            current_path = f"{path}.{key}" if path else key
            
            if key == "soru" and is_valid_question(value):
                questions.append((value, current_path))
            elif is_valid_question(value):
                questions.append((value, current_path))
            elif isinstance(value, dict):
                process_dict(value, current_path)
            elif isinstance(value, list):
                process_list(value, current_path)
    
    def process_list(lst, path=""):
        """Liste tipindeki veriyi işler"""
        for i, item in enumerate(lst):
            current_path = f"{path}[{i}]"
            
            if is_valid_question(item):
                questions.append((item, current_path))
            elif isinstance(item, dict):
                process_dict(item, current_path)
            elif isinstance(item, list):
                process_list(item, current_path)
    
    # Ana veri yapısının tipine göre işlem yap
    try:
        if isinstance(data, dict):
            process_dict(data)
        elif isinstance(data, list):
            process_list(data)
        
        if not questions:
            logger.log_error(file_name, "Soru Bulunamadı", "Veri içinde geçerli soru formatında veri bulunamadı")
            
        # Yol bilgisini logla ve sadece soruları
            
        # Yol bilgisini logla ve sadece soruları döndür
        for i, (_, path) in enumerate(questions):
            logger.log_info(file_name, f"Soru {i+1} yolu: {path}")
            
        return [question for question, _ in questions]
    except Exception as e:
        logger.log_error(file_name, "Soru Çıkarma Hatası", str(e))
        return []

# Ana format düzeltme fonksiyonu
def format_question(item, file_name, logger, soru_index=None):
    """Soru formatını standart formata dönüştürür"""
    question = standard_format.copy()
    
    # Mevcut verileri yeni formata aktar
    for key in item.keys():
        if key in question:
            # Görsel özel işlem - artık basit string formatında bekleniyor
            if key == "gorsel":
                if isinstance(item[key], dict):
                    # Eğer hala eskisi gibi kompleks formatta gelirse (image.data formatı)
                    if "image" in item[key] and "data" in item[key]["image"]:
                        question[key] = item[key]["image"]["data"]
                    else:
                        # Diğer dict formatları
                        logger.log_warning(file_name, "Görsel Format", "Beklenmeyen görsel formatı, boş bırakılıyor", soru_index)
                        question[key] = ""
                else:
                    # Zaten string formatındaysa (doğrudan base64 kodu)
                    question[key] = str(item[key]) if item[key] else ""
                
                # Görsel içeriği kontrol et
                if question[key] and not is_base64_image(question[key]):
                    logger.log_warning(file_name, "Görsel Formatı", "Görsel base64 formatında değil veya bozuk olabilir", soru_index)
            
            # Metin alanlarında HTML etiketlerini koru
            elif key in ["soruMetni", "ustMetin", "cozum"]:
                if item[key] and isinstance(item[key], str):
                    question[key] = item[key]
                else:
                    question[key] = str(item[key]) if item[key] is not None else ""
            
            # Secenekler özel işlemi
            elif key == "secenekler":
                if isinstance(item[key], dict):
                    question[key] = {}
                    for option_key, option_value in item[key].items():
                        if isinstance(option_value, str):
                            # Base64 görsel olup olmadığını kontrol et
                            if is_base64_image(option_value):
                                question[key][option_key] = option_value
                            else:
                                # HTML etiketlerini koru
                                question[key][option_key] = option_value
                        else:
                            # String olmayan değerleri stringe çevir
                            question[key][option_key] = str(option_value)
                else:
                    # Dictionary değilse boş dictionary ata
                    question[key] = {}
                    logger.log_warning(file_name, "Veri Tipi", "Secenekler bir sözlük değil, boş sözlük oluşturuldu", soru_index)
            
            # Kazanım özel işlemi (string ise listeye çevir)
            elif key == "kazanim":
                if isinstance(item[key], str):
                    question[key] = [item[key]]
                elif isinstance(item[key], list):
                    question[key] = item[key]
                else:
                    question[key] = []
            
            # Zorluk seviyesi kontrolü
            elif key == "zorluk":
                try:
                    question[key] = float(item[key])
                    if not (0 <= question[key] <= 1):
                        logger.log_warning(file_name, "Değer Aralığı", f"Zorluk değeri 0-1 aralığında değil: {question[key]}", soru_index)
                except:
                    question[key] = 0
                    logger.log_warning(file_name, "Veri Tipi", "Zorluk değeri sayısal değil", soru_index)
            
            # Diğer alanlar
            else:
                question[key] = item[key]
        else:
            # Standart formatta olmayan alanlar
            logger.log_info(file_name, f"Bilinmeyen alan görmezden gelindi: {key}")
    
    # Gerekli alanların kontrolü ve düzeltmesi
    required_fields = ["soruMetni", "secenekler", "dogruCevap"]
    for field in required_fields:
        if not question[field]:
            logger.log_warning(file_name, "Eksik Alan", f"Gerekli alan boş: {field}", soru_index)
            
            # Boş alanlar için varsayılan değer ekle
            if field == "secenekler":
                question[field] = {}
            else:
                question[field] = ""
    
    return question

# JSON dosyasını işleme fonksiyonu
def process_json_file(input_file, output_file, logger):
    try:
        file_name = os.path.basename(input_file)
        
        # JSON dosyasını oku
        with open(input_file, 'r', encoding='utf-8') as f:
            content = f.read()
            
        # BOM karakteri temizliği
        if content.startswith('\ufeff'):
            content = content[1:]
            logger.log_info(file_name, "BOM karakteri temizlendi")
        
        # JSON içeriğindeki yaygın hataları düzelt ve görsel formatını düzelt
        content = fix_json_content(content)
        
        try:
            # JSON'ı ayrıştır
            data = json.loads(content)
        except json.JSONDecodeError as e:
            error_msg = f"JSON ayrıştırma hatası: Satır {e.lineno}, Kolon {e.colno}"
            logger.log_error(file_name, "JSON Ayrıştırma", error_msg)
            return None
        
        # Tek bir soru objesi mi kontrol et
        if isinstance(data, dict) and 'soruMetni' in data and 'secenekler' in data and 'dogruCevap' in data:
            extracted_questions = [data]
            logger.log_info(file_name, "Tek soru tespit edildi")
        else:
            # Kompleks JSON yapısından soruları çıkar
            extracted_questions = extract_questions(data, file_name, logger)
            
        if not extracted_questions:
            logger.log_error(file_name, "Veri Çıkarma", "Dosyada işlenebilir soru bulunamadı")
            return None
            
        logger.log_info(file_name, f"Toplam {len(extracted_questions)} soru çıkarıldı")
        
        # Soruları standart formata dönüştür
        fixed_data = [format_question(q, file_name, logger, i+1) for i, q in enumerate(extracted_questions)]
        
        # Düzeltilmiş JSON'ı yaz
        with open(output_file, 'w', encoding='utf-8') as f:
            json.dump(fixed_data, f, ensure_ascii=False, indent=2)
        
        logger.log_info(file_name, f"Başarıyla dönüştürüldü: {len(fixed_data)} soru")
        return fixed_data
        
    except Exception as e:
        logger.log_error(file_name, "Dosya İşleme", f"Beklenmeyen hata: {str(e)}")
        traceback.print_exc()
        return None

# Düzeltilmiş JSON'ı Word'e dönüştürme
def json_to_word_profesyonel(json_data, output_file, file_name, logger):
    """
    Düzeltilmiş JSON verilerini profesyonel formatta Word belgesine dönüştürür
    (LGS soru bankası formatında, iki sütunlu, kapak sayfası, cevap anahtarı ile)
    """
    try:
        if not json_data or len(json_data) == 0:
            logger.log_error(file_name, "Word Dönüşüm", "Dönüştürülecek veri yok")
            return False
            
        # Word belgesi oluştur
        doc = Document()
        
        # Sayfa yapısını ayarla (A4 kağıt)
        section = doc.sections[0]
        section.page_width = Inches(8.27)  # A4 genişlik
        section.page_height = Inches(11.69)  # A4 yükseklik
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        
        # İlk sorudan ders bilgisini al
        ders_adi = json_data[0]["ders"] if json_data[0]["ders"] else "LGS Soru Bankası"
        
        # KAPAK SAYFASI OLUŞTUR
        # ---------------------
        # Ders renk kodunu belirle
        if ders_adi == "Türkçe":
            renk_rgb = (204, 0, 0)  # Kırmızı
        elif ders_adi == "Matematik":
            renk_rgb = (0, 51, 153)  # Mavi
        elif ders_adi == "Fen Bilimleri":
            renk_rgb = (0, 153, 51)  # Yeşil
        elif ders_adi == "İnkılap":
            renk_rgb = (204, 102, 0)  # Turuncu
        elif ders_adi == "Din Kültürü":
            renk_rgb = (102, 0, 102)  # Mor
        elif ders_adi == "İngilizce":
            renk_rgb = (153, 153, 0)  # Sarı
        else:
            renk_rgb = (100, 100, 100)  # Gri
            
        # Üst başlık - renkli banner efekti için tablo oluştur
        banner_table = doc.add_table(rows=1, cols=1)
        banner_table.style = 'Table Grid'
        banner_cell = banner_table.cell(0, 0)
        
        # Banner arka plan rengini ayarla
        shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{rgb_to_hex(renk_rgb)}"/>')
        banner_cell._tc.get_or_add_tcPr().append(shading_elm)
        
        # Banner içeriği
        banner_para = banner_cell.paragraphs[0]
        banner_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        banner_para.space_before = Pt(10)
        banner_para.space_after = Pt(10)
        banner_run = banner_para.add_run("LGS SORU BANKASI")
        banner_run.font.size = Pt(36)
        banner_run.font.bold = True
        banner_run.font.color.rgb = RGBColor(255, 255, 255)  # Beyaz yazı
        
        # Alt başlık ekle
        banner_para.add_run("\n")
        banner_alt = banner_para.add_run(f"{ders_adi.upper()} - {datetime.now().year}")
        banner_alt.font.size = Pt(24)
        banner_alt.font.color.rgb = RGBColor(255, 255, 255)  # Beyaz yazı
        
        # Kapak sayfası ortası - Test başlığı
        doc.add_paragraph()
        doc.add_paragraph()
        doc.add_paragraph()
        
        test_baslik = doc.add_paragraph()
        test_baslik.alignment = WD_ALIGN_PARAGRAPH.CENTER
        test_run = test_baslik.add_run("TEST 1")
        test_run.font.size = Pt(72)
        test_run.font.bold = True
        
        # Ders adı
        doc.add_paragraph()
        ders_para = doc.add_paragraph()
        ders_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        ders_run = ders_para.add_run(ders_adi)
        ders_run.font.size = Pt(48)
        ders_run.font.bold = True
        ders_run.font.color.rgb = RGBColor(*renk_rgb)
        
        # Alt bilgi
        doc.add_paragraph()
        doc.add_paragraph()
        doc.add_paragraph()
        doc.add_paragraph()
        
        alt_bilgi = doc.add_paragraph()
        alt_bilgi.alignment = WD_ALIGN_PARAGRAPH.CENTER
        alt_run = alt_bilgi.add_run("SOLOROTA YAYINLARI")
        alt_run.font.size = Pt(12)
        alt_run.font.color.rgb = RGBColor(128, 128, 128)  # Gri
        
        copyright_para = doc.add_paragraph()
        copyright_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        copyright_run = copyright_para.add_run(f"© {datetime.now().year} Tüm Hakları Saklıdır")
        copyright_run.font.size = Pt(10)
        copyright_run.font.color.rgb = RGBColor(128, 128, 128)  # Gri
        
        # Sayfa sonu ekle
        doc.add_page_break()
        
        # İKİ SÜTUNLU SORULAR SAYFASI
        # ---------------------------
        
        # İki sütunlu sayfa yapısı
        section = doc.add_section(WD_SECTION_START.NEW_PAGE)
        section.page_width = Inches(8.27)  # A4 genişlik
        section.page_height = Inches(11.69)  # A4 yükseklik
        sectPr = section._sectPr
        cols = sectPr.xpath('./w:cols')[0]
        cols.set(qn('w:num'), '2')  # 2 sütun
        
        # Sayfa başlığı
        header = section.header
        header_para = header.paragraphs[0]
        header_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        header_run = header_para.add_run(f"TEST 1 - {ders_adi} - {datetime.now().year}")
        header_run.font.size = Pt(10)
        header_run.font.bold = True
        header_run.font.color.rgb = RGBColor(*renk_rgb)
        
        # Filigran efekti için SOLOROTA ekleme (watermark gibi)
        # Not: Word'de gerçek filigran ekleme programatik olarak karmaşık
        filigran_para = doc.add_paragraph()
        filigran_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        filigran_run = filigran_para.add_run("SOLOROTA")
        filigran_run.font.size = Pt(72)
        filigran_run.font.color.rgb = RGBColor(240, 240, 240)  # Çok açık gri
        
        # Her bir soruyu işle ve iki sütuna dağıt
        # Sorular sıralı verilecek
        for soru_index, soru in enumerate(json_data, 1):
            try:
                # Soru paragrafı
                soru_para = doc.add_paragraph()
                
                # Soru numarası - kalın yazılacak
                soru_no_run = soru_para.add_run(f"{soru_index}. ")
                soru_no_run.font.bold = True
                soru_no_run.font.size = Pt(10)
                
                # Üst metin varsa ekle
                if soru["ustMetni"]:
                    ust_metin_run = soru_para.add_run(soru["ustMetni"] + "\n")
                    ust_metin_run.font.size = Pt(9)
                
                # Görsel varsa ekle
                if soru["gorsel"]:
                    # Base64 görsel verisini işle
                    image_data = process_base64_image(soru["gorsel"])
                    if image_data:
                        # Görseli geçici dosyaya kaydet ve ekle
                        temp_image_path = os.path.join(os.path.dirname(output_file), f"temp_img_{soru_index}.png")
                        with open(temp_image_path, "wb") as img_file:
                            img_file.write(image_data)
                        
                        # Görseli ekle - genişliği sütuna sığacak şekilde ayarla
                        soru_para.add_run("\n")
                        max_width = Inches(2.5)  # Sütun genişliğinin biraz altında
                        soru_para.add_picture(temp_image_path, width=max_width)
                        soru_para.add_run("\n")
                        
                        # Geçici dosyayı temizle
                        try:
                            os.remove(temp_image_path)
                        except:
                            pass
                
                # Soru metni
                if soru["soruMetni"]:
                    soru_metni_run = soru_para.add_run("\n" + soru["soruMetni"] + "\n")
                    soru_metni_run.font.size = Pt(9)
                
                # Şıklar
                if soru["secenekler"]:
                    secenekler = soru["secenekler"]
                    for sik, metin in secenekler.items():
                        sik_para = doc.add_paragraph()
                        sik_para.paragraph_format.left_indent = Inches(0.2)
                        
                        # Şık harfini kalın yaz
                        sik_run = sik_para.add_run(f"{sik}) ")
                        sik_run.font.bold = True
                        sik_run.font.size = Pt(9)
                        
                        # Şık değeri base64 görsel mi kontrol et
                        if is_base64_image(metin):
                            # Görsel içeren şık
                            image_data = process_base64_image(metin)
                            if image_data:
                                temp_image_path = os.path.join(os.path.dirname(output_file), f"temp_sik_{soru_index}_{sik}.png")
                                with open(temp_image_path, "wb") as img_file:
                                    img_file.write(image_data)
                                
                                # Şık görselini daha küçük ekle
                                max_width = Inches(1.5)
                                sik_para.add_picture(temp_image_path, width=max_width)
                                
                                # Geçici dosyayı temizle
                                try:
                                    os.remove(temp_image_path)
                                except:
                                    pass
                        else:
                            # Normal metin şık
                            sik_metin_run = sik_para.add_run(metin)
                            sik_metin_run.font.size = Pt(9)
                
                # Sorular arasında boşluk
                doc.add_paragraph()
                
            except Exception as e:
                logger.log_error(file_name, "Soru Word Dönüştürme", str(e), soru_index)
                continue
        
        # CEVAP ANAHTARI SAYFASI
        # ----------------------
        
        # Yeni sayfa ekle - tek sütun
        section = doc.add_section(WD_SECTION_START.NEW_PAGE)
        section.page_width = Inches(8.27)  # A4 genişlik
        section.page_height = Inches(11.69)  # A4 yükseklik
        sectPr = section._sectPr
        cols = sectPr.xpath('./w:cols')[0]
        cols.set(qn('w:num'), '1')  # Tek sütun
        
        # Cevap anahtarı başlığı
        cevap_baslik = doc.add_paragraph()
        cevap_baslik.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cevap_baslik_run = cevap_baslik.add_run(f"TEST 1 - {ders_adi} - CEVAP ANAHTARI")
        cevap_baslik_run.font.size = Pt(18)
        cevap_baslik_run.font.bold = True
        cevap_baslik_run.font.color.rgb = RGBColor(*renk_rgb)
        
        # Alt çizgi
        border_para = doc.add_paragraph()
        border_para.paragraph_format.top_border.width = Pt(1)
        border_para.paragraph_format.top_border.color.rgb = RGBColor(*renk_rgb)
        
        # Cevap anahtarı tablosu - 5 sütunlu
        cevap_table = doc.add_table(rows=1, cols=5)
        cevap_table.style = 'Table Grid'
        cevap_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # Kaç satır gerekecek hesapla
        soru_sayisi = len(json_data)
        satir_sayisi = (soru_sayisi + 4) // 5  # Her satırda 5 cevap, yukarı yuvarla
        
        # Yeterli satır ekle
        for _ in range(satir_sayisi - 1):  # İlk satır zaten var
            cevap_table.add_row()
        
        # Tabloyu doldur
        soru_index = 0
        for row in range(satir_sayisi):
            for col in range(5):
                if soru_index < soru_sayisi:
                    soru = json_data[soru_index]
                    
                    # Hücreyi al
                    cell = cevap_table.cell(row, col)
                    
                    # Soru numarası
                    cell_para = cell.paragraphs[0]
                    cell_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    soru_no_run = cell_para.add_run(f"{soru_index + 1}")
                    soru_no_run.font.bold = True
                    soru_no_run.font.size = Pt(12)
                    
                    # Boşluk
                    cell_para.add_run("\n\n")
                    
                    # Doğru cevap - daire içinde
                    # Not: Word'de gerçek daire çizmek zor, bu yüzden sadece renkli harf yapıyoruz
                    dogru_cevap = soru.get("dogruCevap", "").upper()
                    cevap_run = cell_para.add_run(dogru_cevap)
                    cevap_run.font.size = Pt(14)
                    cevap_run.font.bold = True
                    cevap_run.font.color.rgb = RGBColor(*renk_rgb)
                    
                    soru_index += 1
        
        # Alt bilgi
        alt_bilgi = doc.add_paragraph()
        alt_bilgi.alignment = WD_ALIGN_PARAGRAPH.CENTER
        alt_run = alt_bilgi.add_run(f"SOLOROTA YAYINLARI - {datetime.now().year}")
        alt_run.font.size = Pt(10)
        alt_run.font.color.rgb = RGBColor(128, 128, 128)  # Gri
        
        # Word belgesini kaydet
        try:
            doc.save(output_file)
            logger.log_info(file_name, f"Profesyonel Word soru bankası oluşturuldu: {output_file}")
            return True
        except Exception as e:
            logger.log_error(file_name, "Word Kaydetme", str(e))
            return False
            
    except Exception as e:
        logger.log_error(file_name, "Word Dönüştürme", str(e))
        traceback.print_exc()
        return False

# RGB değerini hex koda dönüştüren yardımcı fonksiyon
def rgb_to_hex(rgb):
    """RGB renk değerini hex koda dönüştürür"""
    r, g, b = rgb
    return f"{r:02x}{g:02x}{b:02x}"

# Gerekli importlar
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
from datetime import datetime
import os
import traceback

# json_to_word_converter.py dosyasında mevcut process_single_file fonksiyonunu değiştir
def process_single_file(json_file):
    """Tek bir JSON dosyasını işleyip profesyonel Word formatına dönüştürür"""
    # Logger oluştur
    logger = Logger(log_file)
    
    file_name = os.path.basename(json_file)
    json_output_file = os.path.join(json_output_dir, file_name)
    word_output_file = os.path.join(word_output_dir, file_name.replace('.json', '_profesyonel.docx'))
    
    print(f"İşleniyor: {file_name}")
    
    # JSON formatını düzelt ve içeriği çıkar
    json_data = process_json_file(json_file, json_output_file, logger)
    
    if json_data:
        print(f"- JSON formatına dönüştürüldü: {len(json_data)} soru")
        
        # Profesyonel Word dosyasını oluştur
        word_success = json_to_word_profesyonel(json_data, word_output_file, file_name, logger)
        if word_success:
            print(f"- Profesyonel Word soru bankası oluşturuldu: {word_output_file}")
        else:
            print(f"- Word dosyası oluşturulamadı!")
    else:
        print(f"- JSON formatına dönüştürülemedi!")
    
    # Log dosyasını oluştur
    logger.save_log()
    
    print(f"\nİşlem tamamlandı:")
    print(f"- Dönüşüm raporu: {log_file}")
    print(f"- JSON çıktısı: {json_output_file}")
    print(f"- Word çıktısı: {word_output_file}")

if __name__ == "__main__":
    print("JSON Soru Formatı Düzeltme ve Word Dönüştürücü")
    print("=" * 60)
    print("Bu program, çeşitli JSON formatlarındaki soru verilerini:")
    print("1. Standart formata dönüştürür")
    print("2. Word belgesi olarak kaydeder")
    print("-" * 60)
    print(f"Kaynak klasör: {input_dir}")
    print(f"JSON çıktı klasörü: {json_output_dir}")
    print(f"Word çıktı klasörü: {word_output_dir}")
    print("-" * 60)
    
    # Kullanıcı seçimine göre işlem yapma
    secim = input("Tüm JSON dosyalarını işlemek için 'T', belirli bir dosyayı işlemek için dosya adını girin: ")
    
    if secim.upper() == 'T':
        process_all_files()
    else:
        # Belirli bir dosyayı işlemek
        json_file = os.path.join(input_dir, secim)
        if os.path.exists(json_file):
            process_single_file(json_file)
        else:
            print(f"Hata: {json_file} dosyası bulunamadı!")
# Ana işlem fonksiyonlarını ekleyin
def process_all_files():
    """Tüm JSON dosyalarını işler"""
    logger = Logger(log_file)
    json_files = glob.glob(os.path.join(input_dir, "*.json"))
    
    for json_file in json_files:
        file_name = os.path.basename(json_file)
        json_output_file = os.path.join(json_output_dir, file_name)
        process_json_file(json_file, json_output_file, logger)

def process_single_file(json_file):
    """Tek bir JSON dosyasını işler"""
    logger = Logger(log_file)
    file_name = os.path.basename(json_file)
    json_output_file = os.path.join(json_output_dir, file_name)
    process_json_file(json_file, json_output_file, logger)
                            